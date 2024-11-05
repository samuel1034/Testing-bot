import os
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
import time

# Path to ChromeDriver
chrome_driver_path = "C:\\Program Files (x86)\\chromedriver.exe"
service = Service(chrome_driver_path)

# Initialize the browser
driver = webdriver.Chrome(service=service)

# URL for login and registration
login_url = "https://isi-web.vercel.app/auth"
registration_url = "https://isi-web.vercel.app/auth/register"

# Create a new document for test results
doc = Document()
doc.add_heading('Resultados de Pruebas de Usabilidad - Módulo Web', level=1)

# Set explicit wait time for loading elements
wait_time = 6

# Create a directory for screenshots
screenshot_dir = "screenshots"
os.makedirs(screenshot_dir, exist_ok=True)


# Helper function to take screenshots
def take_screenshot(name):
    screenshot_path = os.path.join(screenshot_dir, f"{name}.png")
    driver.save_screenshot(screenshot_path)
    doc.add_paragraph(f"Captura de pantalla guardada: {screenshot_path}")


# Function to perform login test
def test_login():
    driver.get(login_url)
    time.sleep(wait_time)
    take_screenshot("login_page_loaded")  # Captura inicial en la página de login
    try:
        # Intentar encontrar y usar los elementos de inicio de sesión
        email_field = WebDriverWait(driver, wait_time).until(
            EC.presence_of_element_located((By.XPATH, "//input[@placeholder='usuario@mail.com']"))
        )
        email_field.send_keys("usuario@mail.com")

        next_button = WebDriverWait(driver, wait_time).until(
            EC.element_to_be_clickable((By.XPATH, "//button[@title='Siguiente']"))
        )
        next_button.click()
        take_screenshot("login_next_button_clicked")

        doc.add_paragraph("Prueba de Login: Elemento encontrado y funcionando.")
    except Exception as e:
        take_screenshot("login_error")
        doc.add_paragraph(f"Prueba de Login: Error: {e}")


# Function to perform registration test with retry on redirection
def test_registration():
    driver.get(registration_url)
    time.sleep(wait_time)
    retry_attempts = 2

    for attempt in range(retry_attempts):
        try:
            create_account_link = WebDriverWait(driver, wait_time).until(
                EC.element_to_be_clickable((By.LINK_TEXT, "Crear una cuenta de ISIPay"))
            )
            create_account_link.click()
            time.sleep(3)

            if driver.current_url == login_url:
                take_screenshot("registration_redirected_to_login")
                print("Redirected back to login. Retrying...")
                continue

            email_field = WebDriverWait(driver, wait_time).until(
                EC.presence_of_element_located((By.XPATH, "//input[@placeholder='usuario@mail.com']"))
            )
            email_field.send_keys("usuario@mail.com")
            take_screenshot("registration_email_entered")

            continue_button = WebDriverWait(driver, wait_time).until(
                EC.element_to_be_clickable((By.XPATH, "//button[@title='Continuar']"))
            )
            continue_button.click()
            take_screenshot("registration_continue_clicked")

            doc.add_paragraph("Prueba de Registro: Elemento encontrado y funcionando.")
            break
        except Exception as e:
            take_screenshot("registration_error")
            doc.add_paragraph(f"Prueba de Registro: Error: {e}")
            break


# Function to test "Iniciar con Google" and capture Google login page text
def test_google_login():
    driver.get(login_url)
    time.sleep(wait_time)

    try:
        google_button = WebDriverWait(driver, wait_time).until(
            EC.element_to_be_clickable((By.XPATH, "//div[contains(text(), 'Iniciar con Google')]"))
        )
        google_button.click()
        take_screenshot("google_login_button_clicked")

        main_window = driver.current_window_handle
        for handle in driver.window_handles:
            if handle != main_window:
                driver.switch_to.window(handle)
                break

        text_element = WebDriverWait(driver, wait_time).until(
            EC.presence_of_element_located((By.TAG_NAME, "body"))
        )
        google_auth_text = text_element.text
        take_screenshot("google_auth_page")

        doc.add_paragraph(f"Texto en la ventana de autenticación de Google: {google_auth_text}")

        driver.close()
        driver.switch_to.window(main_window)

    except Exception as e:
        take_screenshot("google_login_error")
        doc.add_paragraph(f"Iniciar con Google: Error - {e}")


# Perform tests
test_login()
test_registration()
test_google_login()

# Save the document with the results
doc.save("Resultados_Prueba_Modulo_Web.docx")
print("Documento de resultados generado exitosamente con capturas de pantalla.")

# Close the browser
driver.quit()
